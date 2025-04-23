from datetime import timedelta, datetime
from celery import shared_task
from django.contrib import messages
from django.core.mail import EmailMessage
from django.http import HttpResponse
from django.shortcuts import redirect
from django.utils import timezone
from django.utils.timezone import now
from cogu.models import SanitaryIncident, HealthRegion, MajorEvent
from cogu.views import get_actions_for_region, get_actions_taken, get_recommendations, get_next_steps, \
    generate_pdf_report, generate_word_report


def send_report_view(request):
    try:
        context = generate_cogu_report_context()
        send_daily_cogu_report_email(context)  # ✅ maintenant accepté
        messages.success(request, "✅ Rapport envoyé avec succès par e-mail.")
    except Exception as e:
        messages.error(request, f"❌ Échec de l'envoi du rapport : {e}")
    return redirect('incident_report')


# envoie des rapports automatises
def send_daily_cogu_report_email(context=None):
    today = timezone.now().date()
    if context is None:
        context = generate_cogu_report_context()
    docx = generate_word_report_file(context)  # retourne un BytesIO

    email = EmailMessage(
        subject=f"Rapport COGU - {today.strftime('%d/%m/%Y')}",
        body="Veuillez trouver ci-joint le rapport journalier COGU.",
        from_email="noreply@cogu.ci",
        to=["ministre.sante@gouv.ci", "cogu@gouv.ci"]
    )

    email.attach(
        f"COGU_{today.strftime('%d-%m-%Y')}.docx",
        docx.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    email.send()


def send_generate_cogu_report(request, *args, **kwargs):
    today = timezone.now().date()
    yesterday = today - timedelta(days=1)

    output_format = request.GET.get('format') or kwargs.get('format', 'pdf')

    context = generate_cogu_report_context(today, yesterday)

    if output_format == 'pdf':
        return generate_pdf_report(context)
    elif output_format == 'word':
        return generate_word_report(context)
    elif output_format == 'send':
        send_daily_cogu_report_email(context)
        return HttpResponse("Rapport envoyé avec succès.")
    else:
        return HttpResponse("Invalid format specified", status=400)


def generate_cogu_report_context(report_date=None):
    """
    Génère le contexte du rapport COGU pour une date donnée.

    :param report_date: str ou date (optionnel). Si c'est une chaîne, le format attendu est "dd/mm/yyyy" ou "dd-mm-yyyy".
                        Si None, la date du jour est utilisée.
    :return: dict
    """
    # Détermine la date du rapport
    if report_date:
        if isinstance(report_date, str):
            # Essaie de parser la date au format "dd/mm/yyyy" ou "dd-mm-yyyy"
            try:
                report_date_obj = datetime.strptime(report_date, "%d/%m/%Y").date()
            except ValueError:
                try:
                    report_date_obj = datetime.strptime(report_date, "%d-%m-%Y").date()
                except ValueError:
                    # Si le format ne correspond à aucun, on utilise la date du jour
                    report_date_obj = timezone.now().date()
        else:
            report_date_obj = report_date
    else:
        report_date_obj = timezone.now().date()

    # On définit "yesterday" pour le rapport précédent
    yesterday = report_date_obj - timedelta(days=1)

    # Queryset pour les incidents du jour choisi
    daily_incidents = SanitaryIncident.objects.filter(
        date_time__date=report_date_obj
    ).select_related('incident_type', 'city__district__region')

    total_incidents = daily_incidents.count()
    validated_incidents = daily_incidents.filter(status='validated').count()
    pending_incidents = daily_incidents.filter(status='pending').count()

    # Incidents résolus hier
    resolved_incidents = SanitaryIncident.objects.filter(
        date_time__date=yesterday,
        status='validated'
    ).count()

    region_data = []
    # Itérer sur toutes les régions de santé
    for region in HealthRegion.objects.all():
        # Sélectionne les incidents liés à la région
        region_incidents = daily_incidents.filter(city__district__region=region)
        incident_types = {}
        for incident in region_incidents:
            name = incident.incident_type.name
            if name not in incident_types:
                incident_types[name] = {'validated': 0, 'pending': 0, 'total': 0}
            incident_types[name]['total'] += 1
            if incident.status == 'validated':
                incident_types[name]['validated'] += 1
            else:
                incident_types[name]['pending'] += 1
        region_data.append({
            'name': region.name,
            'total_incidents': region_incidents.count(),
            'incident_types': incident_types,
            'actions': get_actions_for_region(region.name)
        })

    return {
        'date': report_date_obj.strftime("%d %B %Y"),
        'total_incidents': total_incidents,
        'validated_incidents': validated_incidents,
        'pending_incidents': pending_incidents,
        'resolved_incidents': resolved_incidents,
        'region_data': region_data,
        'actions_taken': get_actions_taken(),
        'recommendations': get_recommendations(),
        'next_steps': get_next_steps(),
    }
# def generate_cogu_report_context():
#     today = timezone.now().date()
#     yesterday = today - timedelta(days=1)
#
#     daily_incidents = SanitaryIncident.objects.filter(
#         date_time__date=today
#     ).select_related('incident_type', 'city__district__region')
#
#     total_incidents = daily_incidents.count()
#     validated_incidents = daily_incidents.filter(status='validated').count()
#     pending_incidents = daily_incidents.filter(status='pending').count()
#
#     resolved_incidents = SanitaryIncident.objects.filter(
#         date_time__date=yesterday,
#         status='validated'
#     ).count()
#
#     region_data = []
#     for region in HealthRegion.objects.all():
#         region_incidents = daily_incidents.filter(city__district__region=region)
#         incident_types = {}
#         for incident in region_incidents:
#             name = incident.incident_type.name
#             if name not in incident_types:
#                 incident_types[name] = {'validated': 0, 'pending': 0, 'total': 0}
#             incident_types[name]['total'] += 1
#             if incident.status == 'validated':
#                 incident_types[name]['validated'] += 1
#             else:
#                 incident_types[name]['pending'] += 1
#         region_data.append({
#             'name': region.name,
#             'total_incidents': region_incidents.count(),
#             'incident_types': incident_types,
#             'actions': get_actions_for_region(region.name)
#         })
#
#     return {
#         'date': today.strftime("%d %B %Y"),
#         'total_incidents': total_incidents,
#         'validated_incidents': validated_incidents,
#         'pending_incidents': pending_incidents,
#         'resolved_incidents': resolved_incidents,
#         'region_data': region_data,
#         'actions_taken': get_actions_taken(),
#         'recommendations': get_recommendations(),
#         'next_steps': get_next_steps(),
#     }


def generate_word_report_file(context):
    from docx import Document
    from io import BytesIO

    document = Document()
    document.add_heading('RAPPORT JOURNALIER COGU', level=0)
    document.add_paragraph(f"Date : {context['date']}")
    document.add_paragraph("Destinataires : Monsieur le Ministre de la Santé, Membres du COGU")
    document.add_paragraph("Émetteur : Directeur Général de la Santé et de l'Hygiène Publique (DGSHP)")

    document.add_heading('RÉCAPITULATIF GLOBAL', level=1)
    document.add_paragraph(f"Nombre total d'incidents signalés : {context['total_incidents']}")
    document.add_paragraph(f"Incidents validés : {context['validated_incidents']}")
    document.add_paragraph(f"Incidents en cours de validation : {context['pending_incidents']}")
    document.add_paragraph(f"Incidents résolus hier : {context['resolved_incidents']}")

    document.add_heading('DÉTAILS PAR RÉGION SANITAIRE', level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Région'
    hdr[1].text = 'Total'
    hdr[2].text = 'Types d’incidents'
    hdr[3].text = 'Statuts'
    hdr[4].text = 'Actions'

    for region in context['region_data']:
        row = table.add_row().cells
        row[0].text = region['name']
        row[1].text = str(region['total_incidents'])
        row[2].text = "\n".join([f"- {v['total']} cas de {k}" for k, v in region['incident_types'].items()])
        row[3].text = "\n".join(
            [f"- {v['validated']} validés, {v['pending']} en cours" for v in region['incident_types'].values()])
        row[4].text = "\n".join([f"- {a}" for a in region['actions']])

    for section, items in [
        ("ACTIONS MENÉES ET INTERVENTIONS EN COURS", context['actions_taken']),
        ("RECOMMANDATIONS ET PERSPECTIVES", context['recommendations']),
        ("PROCHAINES ÉTAPES", context['next_steps']),
    ]:
        document.add_heading(section, level=1)
        for item in items:
            document.add_paragraph(item, style='List Bullet')

    document.add_heading("Conclusion", level=1)
    document.add_paragraph(
        "La situation reste sous contrôle. Les investigations se poursuivent. "
        "Un suivi quotidien est assuré pour informer Monsieur le Ministre."
    )

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


@shared_task
def task_send_daily_cogu_report():
    send_daily_cogu_report_email()


def auto_send_if_major_event():
    today = timezone.now().date()
    if MajorEvent.objects.filter(start_date__date__lte=today, end_date__date__gte=today).exists():
        context = generate_cogu_report_context()
        send_daily_cogu_report_email(context)
