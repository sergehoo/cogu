import base64
import calendar
import csv
import json
import os
import time
from collections import defaultdict
from datetime import timedelta

import folium
from django.conf import settings
from django.contrib import messages
from django.contrib.auth import get_user_model
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.gis.db.models.functions import AsGeoJSON
from django.core.files.base import ContentFile
from django.core.paginator import Paginator
from django.core.serializers import serialize
from django.db import models
from django.db.models import Q, Count, F, Sum, Min
from django.http import JsonResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.template import context
from django.urls import reverse_lazy
from django.utils import timezone
from django.utils.dateparse import parse_date
from django.utils.html import format_html
from django.utils.timezone import localtime
from django.views.decorators.http import require_POST
from django.views.generic import TemplateView, ListView, CreateView, DetailView, UpdateView, DeleteView, FormView
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from selenium import webdriver
from selenium.webdriver.chrome.options import Options

from cogu.filters import PatientFilter
from cogu.forms import SanitaryIncidentForm, PublicIncidentForm, ContactForm, CoguReportForm
from cogu.models import Patient, MajorEvent, IncidentType, SanitaryIncident, Commune, HealthRegion, VictimCare, \
    WhatsAppMessage, DistrictSanitaire, PolesRegionaux, Kit, Fournisseur, Stock, KitCategorie, CoguReport
from django.contrib.gis.geos import Point, GEOSGeometry
from django.http import HttpResponse
from django.template.loader import get_template
from django.shortcuts import render
from django.utils import timezone
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from xhtml2pdf import pisa
from io import BytesIO
from .models import SanitaryIncident, IncidentType, MajorEvent, Commune
import matplotlib

matplotlib.use('Agg')  # Important pour les serveurs/django/macos
import matplotlib.pyplot as plt
import io


# Create your views here.
class RoleRequiredMixin(UserPassesTestMixin):
    allowed_roles = []
    redirect_view_if_denied = 'public_dashboard'  # nom de l’URL

    def test_func(self):
        user = self.request.user
        return user.is_authenticated and user.roleemployee in self.allowed_roles

    def handle_no_permission(self):
        # Redirige vers la vue publique si le test échoue
        return redirect(self.redirect_view_if_denied)


class LandingView(FormView):
    # template_name = "pages/landing.html"
    template_name = "pages/landing.html"
    form_class = ContactForm
    success_url = reverse_lazy('landing')

    def form_valid(self, form):
        form.save()
        messages.success(self.request, "Votre message a été envoyé avec succès. Nous vous contacterons bientôt!")
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, "Il y a eu une erreur dans l'envoi de votre message. Veuillez réessayer.")
        return super().form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Récupération des données pour les statistiques
        # context['form'] = ContactForm()
        context['incidents_count'] = SanitaryIncident.objects.count()
        context['active_cases'] = SanitaryIncident.objects.filter(
            status='validated',
            outcome__in=['mort', 'blessure']
        ).count()
        context['regions_count'] = HealthRegion.objects.count()
        context['interventions_count'] = VictimCare.objects.count()

        # Incidents récents pour la sidebar
        context['recent_incidents'] = SanitaryIncident.objects.select_related(
            'incident_type', 'city'
        ).order_by('-date_time')[:5]

        # Données pour la carte
        context['map_incidents'] = SanitaryIncident.objects.filter(
            location__isnull=False
        ).select_related('incident_type')[:20]

        # Régions sanitaires avec leurs districts
        context['health_regions'] = HealthRegion.objects.annotate(
            district_count=Count('districts')
        ).prefetch_related('districts').order_by('name')

        return context


class PolitiqueConfidentialiteView(TemplateView):
    template_name = "pages/politique_confidential.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Récupération des données pour les statistiques
        context['incidents_count'] = SanitaryIncident.objects.count()
        context['active_cases'] = SanitaryIncident.objects.filter(
            status='validated',
            outcome__in=['mort', 'blessure']
        ).count()
        context['regions_count'] = HealthRegion.objects.count()
        context['interventions_count'] = VictimCare.objects.count()

        # Incidents récents pour la sidebar
        context['recent_incidents'] = SanitaryIncident.objects.select_related(
            'incident_type', 'city'
        ).order_by('-date_time')[:5]

        # Données pour la carte
        context['map_incidents'] = SanitaryIncident.objects.filter(
            location__isnull=False
        ).select_related('incident_type')[:20]

        # Régions sanitaires avec leurs districts
        context['health_regions'] = HealthRegion.objects.annotate(
            district_count=Count('districts')
        ).prefetch_related('districts').order_by('name')

        return context


class PublicUserDashboard(LoginRequiredMixin, TemplateView):
    template_name = "pages/public/public_dashboard.html"
    login_url = 'account_login'

    def get_context_data(self, **kwargs):
        context = super().get_context_data()
        context['mesincidents'] = SanitaryIncident.objects.filter(posted_by=self.request.user).order_by('-created_at')[
                                  :5]
        return context


class PublicIncidentCreateView(LoginRequiredMixin, CreateView):
    model = SanitaryIncident
    template_name = 'pages/public/sanitaryincidentcreate.html'
    form_class = PublicIncidentForm
    success_url = reverse_lazy('public_incidentlist')

    def form_valid(self, form):
        instance = form.save(commit=False)  # Ne pas sauvegarder tout de suite
        User = get_user_model()

        if isinstance(self.request.user, User):
            instance.posted_by = self.request.user
        # instance.posted_by = self.request.user  # Assignation de l'utilisateur
        messages.success(self.request, 'Incident enregistré avec succès!')
        instance.save()
        # form.save_m2m()  # Important si tu as des champs ManyToMany
        return redirect('public_incidentlist')

    def form_invalid(self, form):
        messages.error(self.request, 'Veuillez corriger les erreurs ci-dessous :')

        # Boucle sur les champs pour afficher chaque erreur individuellement
        for field, errors in form.errors.items():
            field_label = form.fields.get(field).label if field in form.fields else field
            for error in errors:
                messages.error(self.request, format_html("<strong>{}</strong>: {}", field_label, error))

        return super().form_invalid(form)


class PublicIncidentListView(LoginRequiredMixin, ListView):
    model = SanitaryIncident
    template_name = 'pages/public/public_incident.html'
    context_object_name = 'incidents'
    paginate_by = 10
    ordering = ['-date_time']

    def get_queryset(self):
        return SanitaryIncident.objects.filter(posted_by=self.request.user).order_by(*self.ordering)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['incidents_active'] = True  # pour activer le menu dans le template
        return context


class PublicIncidentDetailView(LoginRequiredMixin, DetailView):
    model = SanitaryIncident
    template_name = 'pages/public/public_incident_details.html'
    context_object_name = 'incidentsdetails'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['incidents_active'] = True  # pour activer le menu dans le template
        return context


class CADashborad(LoginRequiredMixin, TemplateView):
    template_name = "pages/dashboard.html"
    login_url = 'account_login'
    allowed_roles = ['National', 'Regional']
    redirect_view_if_denied = 'public_dashboard'

    def get(self, request, *args, **kwargs):
        now = timezone.now()
        last_month = now - timedelta(days=30)

        event_id = request.GET.get('event_id')
        incidents_qs = SanitaryIncident.objects.all()
        if event_id and event_id != 'all':
            incidents_qs = incidents_qs.filter(event_id=event_id)

        incidents_count = incidents_qs.count()

        active_cases_qs = incidents_qs.filter(Q(outcome='mort') | Q(outcome='blessure'))
        active_cases_count = active_cases_qs.count()

        interventions_qs = incidents_qs.filter(patients_related__isnull=False).distinct()
        interventions_count = interventions_qs.count()

        resolved_cases_qs = incidents_qs.filter(outcome='mort')
        resolved_cases_count = resolved_cases_qs.count()

        incidents_last_month = incidents_qs.filter(date_time__gte=last_month).count()
        incidents_monthly_change = self.calculate_change(incidents_last_month, incidents_count)

        active_cases_last_month = active_cases_qs.filter(date_time__gte=last_month).count()
        interventions_last_month = interventions_qs.filter(date_time__gte=last_month).count()
        resolved_last_month = resolved_cases_qs.filter(date_time__gte=last_month).count()

        incidents_percentage = (incidents_last_month / max(1, incidents_count)) * 100
        active_cases_percentage = (active_cases_count / max(1, incidents_count)) * 100
        interventions_percentage = (interventions_count / max(1, incidents_count)) * 100
        resolved_percentage = (resolved_cases_count / max(1, incidents_count)) * 100

        incidents_critical = incidents_qs.filter(outcome='mort').order_by('-date_time')[:1]

        recent_incidents = incidents_qs.select_related('incident_type', 'city').order_by('-date_time')
        paginator = Paginator(recent_incidents, 5)
        page = request.GET.get("page")
        recent_incidents_page = paginator.get_page(page)

        # total_people_involved = incidents_qs.aggregate(total=models.Sum('number_of_people_involved'))['total'] or 0
        #
        # outcomes = ['mort', 'evacue', 'pris_charge', 'blessure', 'exeat']
        # stats = {}
        # for o in outcomes:
        #     count = incidents_qs.filter(outcome=o).count()
        #     stats[o] = {
        #         'count': count,
        #         'percent': round((count / max(1, total_people_involved)) * 100, 2)
        #     }

        # Total des personnes impliquées
        total_people_involved = incidents_qs.aggregate(total=Sum('number_of_people_involved'))['total'] or 0

        # Totaux par type
        sums = incidents_qs.aggregate(
            deces_total=Sum('deces_nbr'),
            evacues_total=Sum('evacues_nbr'),
            pris_en_charge_total=Sum('pris_en_charge_nbr'),
            blessure_total=Sum('blessure_nbr'),
            exeat_total=Sum('exeat_nbr'),
        )

        # Pourcentage helper
        def percent(value):
            return round((value / total_people_involved) * 100, 2) if total_people_involved > 0 else 0.0

        # Données agrégées à inclure dans le contexte
        poles_data = self.fetch_poles_data(event_id)
        regions_data = self.get_regions_data(event_id)
        districts_data = self.fetch_districts_data(event_id)
        hierarchical_data = self.get_hierarchical_data(event_id)
        incident_types_data = self.get_incident_types_data(event_id)
        event_distribution = self.get_event_distribution()
        chart_data = self.get_highcharts_data(event_id)
        monthly_data = self.get_monthly_trends()
        event_stats = self.get_event_stats()

        context = {
            'events': MajorEvent.objects.all().order_by('-start_date'),
            'event_id': event_id,

            'people_involved': total_people_involved,
            'deces_count': sums['deces_total'] or 0,
            'deces_percentage': percent(sums['deces_total'] or 0),

            'evacue_count': sums['evacues_total'] or 0,
            'evacue_percentage': percent(sums['evacues_total'] or 0),

            'pris_charge_count': sums['pris_en_charge_total'] or 0,
            'pris_charge_percentage': percent(sums['pris_en_charge_total'] or 0),

            'blessure_count': sums['blessure_total'] or 0,
            'blessure_percentage': percent(sums['blessure_total'] or 0),

            'exeat_count': sums['exeat_total'] or 0,
            'exeat_percentage': percent(sums['exeat_total'] or 0),

            'incidents_count': incidents_count,
            'active_cases': active_cases_count,
            'interventions_count': interventions_count,
            'resolved_cases': resolved_cases_count,
            'incidents_percentage': round(incidents_percentage, 1),
            'incidents_monthly_change': incidents_monthly_change,
            'active_cases_percentage': round(active_cases_percentage, 1),
            'active_cases_change': self.calculate_change(active_cases_last_month, active_cases_count),
            'interventions_percentage': round(interventions_percentage, 1),
            'interventions_change': self.calculate_change(interventions_last_month, interventions_count),
            'resolved_percentage': round(resolved_percentage, 1),
            'resolved_change': self.calculate_change(resolved_last_month, resolved_cases_count),

            'incidents_critical': incidents_critical,
            'recent_incidents': recent_incidents_page,
            'incident_types': IncidentType.objects.all(),

            'poles_labels': [p['name'] for p in poles_data],
            'poles_data': [p['count'] for p in poles_data],
            'regions_labels': [r['name'] for r in regions_data],
            'regions_data': [r['count'] for r in regions_data],
            'districts_labels': [d['name'] for d in districts_data],
            'districts_data': [d['count'] for d in districts_data],
            'incident_types_labels': [t['name'] for t in incident_types_data],
            'incident_types_data': [t['count'] for t in incident_types_data],
            'hierarchical_data': hierarchical_data,
            'available_years': list(range(2019, timezone.now().year + 1)),
            'event_labels': [e['name'] for e in event_distribution],
            'event_data': [e['count'] for e in event_distribution],
            "highcharts_categories": chart_data["categories"],
            "highcharts_series": chart_data["series"],
            "highcharts_outcome_pie": chart_data["pie_data"],
            'event_stats': event_stats,
            'monthly_labels': monthly_data['labels'],
            'monthly_current_data': monthly_data['current_year'],
            'monthly_previous_data': monthly_data['previous_year'],
            'monthly_current_label': monthly_data['current_year_label'],
            'monthly_previous_label': monthly_data['previous_year_label'],
        }

        return self.render_to_response(context)

    # def get(self, request, *args, **kwargs):
    #     now = timezone.now()
    #     last_month = now - timedelta(days=30)
    #
    #     selected_event_id = request.GET.get('event_id')
    #     incidents_qs = SanitaryIncident.objects.all()
    #
    #     if selected_event_id and selected_event_id != 'all':
    #         incidents_qs = incidents_qs.filter(event__id=selected_event_id)
    #
    #     incidents_count = incidents_qs.count()
    #
    #     active_cases_qs = incidents_qs.filter(Q(outcome='mort') | Q(outcome='blessure'))
    #     active_cases_count = active_cases_qs.count()
    #
    #     interventions_qs = incidents_qs.filter(patients_related__isnull=False).distinct()
    #     interventions_count = interventions_qs.count()
    #
    #     resolved_cases_qs = incidents_qs.filter(outcome='mort')
    #     resolved_cases_count = resolved_cases_qs.count()
    #
    #     incidents_last_month = incidents_qs.filter(date_time__gte=last_month).count()
    #     active_cases_last_month = active_cases_qs.filter(date_time__gte=last_month).count()
    #     interventions_last_month = interventions_qs.filter(date_time__gte=last_month).count()
    #     resolved_last_month = resolved_cases_qs.filter(date_time__gte=last_month).count()
    #
    #     incidents_percentage = (incidents_last_month / max(1, incidents_count)) * 100
    #     active_cases_percentage = (active_cases_count / max(1, incidents_count)) * 100
    #     interventions_percentage = (interventions_count / max(1, incidents_count)) * 100
    #     resolved_percentage = (resolved_cases_count / max(1, incidents_count)) * 100
    #
    #     incidents_critical = incidents_qs.filter(outcome='mort').order_by('-date_time')[:1]
    #
    #     recent_incidents = SanitaryIncident.objects.select_related('incident_type', 'city').order_by('-date_time')
    #     paginator = Paginator(recent_incidents, 5)
    #     page = request.GET.get("page")
    #     recent_incidents_page = paginator.get_page(page)
    #
    #     poles_data = self.fetch_poles_data()
    #     districts_data = self.fetch_districts_data()
    #     incident_types_data = self.get_incident_types_data()
    #     regions_data = self.get_regions_data()
    #     monthly_data = self.get_monthly_trends()
    #     event_distribution = self.get_event_distribution()
    #
    #     incidents = incidents_qs
    #     total_incidents = incidents.count()
    #     total_people_involved = incidents.aggregate(total=models.Sum('number_of_people_involved'))['total'] or 0
    #
    #     outcomes = ['mort', 'evacue', 'pris_charge', 'blessure', 'exeat']
    #     stats = {}
    #     for o in outcomes:
    #         count = incidents.filter(outcome=o).count()
    #         stats[o] = {
    #             'count': count,
    #             'percent': round((count / total_incidents) * 100, 2) if total_incidents else 0
    #         }
    #
    #     event_stats = self.get_event_stats()
    #     hierarchical_data = self.get_hierarchical_data()
    #     chart_data = self.get_highcharts_data()
    #
    #     context = {
    #         'events': MajorEvent.objects.all().order_by('-start_date'),
    #         'event_id': selected_event_id,
    #         'people_involved': total_people_involved,
    #         'deces_count': stats['mort']['count'],
    #         'deces_percentage': stats['mort']['percent'],
    #         'evacue_count': stats['evacue']['count'],
    #         'evacue_percentage': stats['evacue']['percent'],
    #         'pris_charge_count': stats['pris_charge']['count'],
    #         'pris_charge_percentage': stats['pris_charge']['percent'],
    #         'blessure_count': stats['blessure']['count'],
    #         'blessure_percentage': stats['blessure']['percent'],
    #         'exeat_count': stats['exeat']['count'],
    #         'exeat_percentage': stats['exeat']['percent'],
    #         "highcharts_categories": chart_data["categories"],
    #         "highcharts_series": chart_data["series"],
    #         "highcharts_outcome_pie": chart_data["pie_data"],
    #         'event_stats': event_stats,
    #         'hierarchical_data': hierarchical_data,
    #         'available_years': list(range(2019, timezone.now().year + 1)),
    #         'event_labels': [e['name'] for e in event_distribution],
    #         'event_data': [e['count'] for e in event_distribution],
    #         'incidents_count': incidents_count,
    #         'active_cases': active_cases_count,
    #         'interventions_count': interventions_count,
    #         'resolved_cases': resolved_cases_count,
    #         'incidents_percentage': round(incidents_percentage, 1),
    #         'incidents_monthly_change': self.calculate_change(incidents_last_month, incidents_count),
    #         'active_cases_percentage': round(active_cases_percentage, 1),
    #         'active_cases_change': self.calculate_change(active_cases_last_month, active_cases_count),
    #         'interventions_percentage': round(interventions_percentage, 1),
    #         'interventions_change': self.calculate_change(interventions_last_month, interventions_count),
    #         'resolved_percentage': round(resolved_percentage, 1),
    #         'resolved_change': self.calculate_change(resolved_last_month, resolved_cases_count),
    #         'incidents_critical': incidents_critical,
    #         'recent_incidents': recent_incidents_page,
    #         'incident_types': IncidentType.objects.all(),
    #         'regions_labels': [r['name'] for r in regions_data],
    #         'regions_data': [r['count'] for r in regions_data],
    #         'monthly_labels': monthly_data['labels'],
    #         'monthly_current_data': monthly_data['current_year'],
    #         'monthly_previous_data': monthly_data['previous_year'],
    #         'monthly_current_label': monthly_data['current_year_label'],
    #         'monthly_previous_label': monthly_data['previous_year_label'],
    #         'incident_types_labels': [t['name'] for t in incident_types_data],
    #         'incident_types_data': [t['count'] for t in incident_types_data],
    #         'poles_labels': [p['name'] for p in poles_data],
    #         'poles_data': [p['count'] for p in poles_data],
    #         'districts_labels': [d['name'] for d in districts_data],
    #         'districts_data': [d['count'] for d in districts_data],
    #     }
    #
    #     return self.render_to_response(context)

    def get_hierarchical_data(self, event_id=None):
        from collections import defaultdict

        data = defaultdict(lambda: {
            'count': 0,
            'regions': defaultdict(lambda: {
                'count': 0,
                'types': defaultdict(int)
            })
        })

        qs = SanitaryIncident.objects.select_related('incident_type', 'city__district__region__poles')
        if event_id and event_id != 'all':
            qs = qs.filter(event_id=event_id)

        for incident in qs:
            try:
                pole_name = incident.city.district.region.poles.name
                region_name = incident.city.district.region.name
                incident_type_name = incident.incident_type.name
            except AttributeError:
                continue

            data[pole_name]['count'] += 1
            data[pole_name]['regions'][region_name]['count'] += 1
            data[pole_name]['regions'][region_name]['types'][incident_type_name] += 1

        poles_list = []
        for pole_name, pole_data in data.items():
            regions_list = []
            for region_name, region_data in pole_data['regions'].items():
                types_list = sorted([
                    {'name': tname, 'count': tcount} for tname, tcount in region_data['types'].items()
                ], key=lambda x: x['count'], reverse=True)
                regions_list.append({
                    'name': region_name,
                    'count': region_data['count'],
                    'types': types_list
                })

            regions_list = sorted(regions_list, key=lambda x: x['count'], reverse=True)
            poles_list.append({
                'name': pole_name,
                'count': pole_data['count'],
                'regions': regions_list
            })

        return sorted(poles_list, key=lambda x: x['count'], reverse=True)

    def calculate_change(self, last, current):
        if last == 0:
            return current * 100 if current else 0
        return round(((current - last) / last) * 100, 1)

    def get_highcharts_data(self, event_id=None):
        from collections import defaultdict
        from django.db.models.functions import TruncDate

        series_by_event = defaultdict(lambda: defaultdict(int))
        categories_set = set()

        qs = SanitaryIncident.objects.filter(event__isnull=False).annotate(date=TruncDate("date_time"))
        if event_id and event_id != 'all':
            qs = qs.filter(event_id=event_id)

        for row in qs.values("event__name", "incident_type__name", "date").annotate(count=Count("id")):
            event = row["event__name"]
            inc_type = row["incident_type__name"]
            date = row["date"].strftime("%Y-%m-%d")
            categories_set.add(date)
            key = f"{event} | {inc_type}"
            series_by_event[key][date] += row["count"]

        categories = sorted(list(categories_set))
        series_data = []
        for label, data in series_by_event.items():
            series_data.append({
                "name": label,
                "data": [data.get(date, 0) for date in categories]
            })

        outcome_data = SanitaryIncident.objects.values("outcome").annotate(count=Count("id")).order_by("-count")
        pie_data = [{"name": row["outcome"].capitalize(), "y": row["count"]} for row in outcome_data]

        return {
            "categories": categories,
            "series": series_data,
            "pie_data": pie_data,
        }

    def get_event_stats(self):
        from collections import defaultdict

        stats = defaultdict(lambda: defaultdict(int))

        incidents = SanitaryIncident.objects.select_related('event', 'incident_type').filter(event__isnull=False)

        for incident in incidents:
            event_name = incident.event.name if incident.event else "Inconnu"
            incident_type_name = incident.incident_type.name
            stats[event_name][incident_type_name] += 1

        # Format final : trié par nombre total d’incidents
        result = []
        for event_name, types in stats.items():
            types_list = [{'name': k, 'count': v} for k, v in types.items()]
            types_list = sorted(types_list, key=lambda x: x['count'], reverse=True)
            result.append({
                'name': event_name,
                'count': sum(t['count'] for t in types_list),
                'types': types_list
            })

        return sorted(result, key=lambda x: x['count'], reverse=True)

    def get_event_distribution(self):
        return SanitaryIncident.objects.filter(event__isnull=False).values(name=F('event__name')).annotate(
            count=Count('id')).order_by('-count')

    def get_monthly_trends(self):
        from collections import defaultdict

        current_year = timezone.now().year
        previous_year = current_year - 1
        MONTHS_FR = ["", "Janvier", "Février", "Mars", "Avril", "Mai", "Juin", "Juillet", "Août", "Septembre",
                     "Octobre", "Novembre", "Décembre"]

        current_data = defaultdict(int)
        previous_data = defaultdict(int)

        incidents = SanitaryIncident.objects.filter(date_time__year__in=[current_year, previous_year]).annotate(
            month=F('date_time__month'), year=F('date_time__year')).values('month', 'year').annotate(count=Count('id'))

        for entry in incidents:
            if entry['year'] == current_year:
                current_data[entry['month']] = entry['count']
            elif entry['year'] == previous_year:
                previous_data[entry['month']] = entry['count']

        monthly_labels = [f"{MONTHS_FR[m]}" for m in range(1, 13)]
        current_year_data = [current_data[m] for m in range(1, 13)]
        previous_year_data = [previous_data[m] for m in range(1, 13)]

        return {
            'labels': monthly_labels,
            'current_year': current_year_data,
            'previous_year': previous_year_data,
            'current_year_label': current_year,
            'previous_year_label': previous_year
        }

    def get_incident_types_data(self, event_id=None):
        qs = SanitaryIncident.objects.all()
        if event_id and event_id != 'all':
            qs = qs.filter(event_id=event_id)
        return qs.values(name=F('incident_type__name')).annotate(count=Count('id')).order_by('-count')

    def fetch_poles_data(self, event_id=None):
        qs = SanitaryIncident.objects.filter(city__district__region__poles__isnull=False)
        if event_id and event_id != 'all':
            qs = qs.filter(event_id=event_id)
        return qs.values(name=F('city__district__region__poles__name')).annotate(count=Count('id')).order_by('-count')

    def fetch_districts_data(self, event_id=None):
        qs = SanitaryIncident.objects.filter(city__district__isnull=False)
        if event_id and event_id != 'all':
            qs = qs.filter(event_id=event_id)
        return qs.values(name=F('city__district__nom')).annotate(count=Count('id')).order_by('-count')

    def get_regions_data(self, event_id=None):
        qs = SanitaryIncident.objects.filter(city__district__region__isnull=False)
        if event_id and event_id != 'all':
            qs = qs.filter(event_id=event_id)
        return qs.values(name=F('city__district__region__name')).annotate(count=Count('id')).order_by('-count')

    def get_event_distribution(self):
        return (
            SanitaryIncident.objects
            .filter(event__isnull=False)
            .values(name=F('event__name'))
            .annotate(count=Count('id'))
            .order_by('-count')
        )

    def get_incident_types_data(self, event_id=None):
        qs = SanitaryIncident.objects.all()
        if event_id and event_id != 'all':
            qs = qs.filter(event_id=event_id)
        return qs.values(name=F('incident_type__name')).annotate(count=Count('id')).order_by('-count')


class PatientListView(LoginRequiredMixin, RoleRequiredMixin, ListView):
    model = Patient
    template_name = 'patient/list.html'
    context_object_name = 'patients'
    allowed_roles = ['National', 'Regional']
    paginate_by = 20
    ordering = ['-created_at']
    filterset_class = PatientFilter

    def test_func(self):
        return self.request.user.is_authenticated and self.request.user.roleemployee in ['National', 'Regional']

    def get_queryset(self):
        queryset = super().get_queryset()
        search_query = self.request.GET.get('search')

        if search_query:
            queryset = queryset.filter(
                Q(nom__icontains=search_query) |
                Q(prenoms__icontains=search_query) |
                Q(code_patient__icontains=search_query) |
                Q(contact__icontains=search_query)
            )

        return queryset


class PatientCreateView(LoginRequiredMixin, RoleRequiredMixin, CreateView):
    model = Patient
    template_name = 'patient/create.html'
    allowed_roles = ['National', 'Regional']
    fields = '__all__'
    success_url = reverse_lazy('patient_list')

    def form_valid(self, form):
        form.instance.created_by = self.request.user
        response = super().form_valid(form)
        messages.success(self.request, f"Patient {self.object} créé avec succès!")
        return response

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = "Nouveau Patient"
        return context


class PatientDetailView(LoginRequiredMixin, RoleRequiredMixin, DetailView):
    model = Patient
    template_name = 'patient/detail.html'
    allowed_roles = ['National', 'Regional']
    context_object_name = 'patient'
    slug_field = 'code_patient'
    slug_url_kwarg = 'code_patient'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f"Détails Patient - {self.object.code_patient}"
        return context


class PatientUpdateView(LoginRequiredMixin, RoleRequiredMixin, UpdateView):
    model = Patient
    template_name = 'patient/update.html'
    fields = '__all__'
    slug_field = 'code_patient'
    slug_url_kwarg = 'code_patient'
    success_url = reverse_lazy('patient_list')
    allowed_roles = ['National', 'Regional']

    def form_valid(self, form):
        response = super().form_valid(form)
        messages.success(self.request, f"Patient {self.object} mis à jour avec succès!")
        return response

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = f"Modifier Patient - {self.object.code_patient}"
        return context


class PatientDeleteView(LoginRequiredMixin, RoleRequiredMixin, DeleteView):
    model = Patient
    template_name = 'patient/delete.html'
    slug_field = 'code_patient'
    slug_url_kwarg = 'code_patient'
    success_url = reverse_lazy('patient_list')
    allowed_roles = ['National', 'Regional']

    def delete(self, request, *args, **kwargs):
        response = super().delete(request, *args, **kwargs)
        messages.success(request, f"Patient {self.object} supprimé avec succès!")
        return response


class MajorEventListView(LoginRequiredMixin, RoleRequiredMixin, ListView):
    model = MajorEvent
    template_name = 'majorevent/list.html'
    context_object_name = 'events'
    paginate_by = 20
    ordering = ['-start_date']
    allowed_roles = ['National', 'Regional']


class MajorEventGridView(LoginRequiredMixin, RoleRequiredMixin, ListView):
    model = MajorEvent
    template_name = 'majorevent/event_grid.html'
    context_object_name = 'events'
    paginate_by = 20
    # ordering = ['-start_date']
    allowed_roles = ['National', 'Regional']

    # def get_queryset(self):
    #     now = timezone.now()
    #     return MajorEvent.objects.filter(start_date__gte=now).order_by('start_date')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        for event in context['events']:
            event.mort_count = event.sanitaryincident_set.filter(outcome='mort').count()
            event.blessure_count = event.sanitaryincident_set.filter(outcome='blessure').count()
            event.sauve_count = event.sanitaryincident_set.filter(outcome='sauvé').count()
            event.incident_count = event.sanitaryincident_set.all().count()
        return context


class MajorEventCreateView(LoginRequiredMixin, RoleRequiredMixin, CreateView):
    model = MajorEvent
    template_name = 'majorevent/event_create.html'
    fields = '__all__'
    success_url = reverse_lazy('majorevent_list')
    allowed_roles = ['National', 'Regional']


class MajorEventDetailView(LoginRequiredMixin, RoleRequiredMixin, DetailView):
    model = MajorEvent
    template_name = 'majorevent/event_detail.html'
    context_object_name = 'event'
    allowed_roles = ['National', 'Regional']


class MajorEventUpdateView(LoginRequiredMixin, RoleRequiredMixin, UpdateView):
    model = MajorEvent
    template_name = 'majorevent/event_update.html'
    fields = '__all__'
    success_url = reverse_lazy('majorevent_list')
    allowed_roles = ['National', 'Regional']


class MajorEventDeleteView(LoginRequiredMixin, RoleRequiredMixin, DeleteView):
    model = MajorEvent
    template_name = 'majorevent/event_delete.html'
    success_url = reverse_lazy('majorevent_list')
    allowed_roles = ['National', 'Regional']


class IncidentTypeListView(LoginRequiredMixin, RoleRequiredMixin, ListView):
    model = IncidentType
    template_name = 'incidenttype/incidentlist.html'
    context_object_name = 'types'
    ordering = ['name']
    allowed_roles = ['National', 'Regional']


class IncidentTypeCreateView(LoginRequiredMixin, RoleRequiredMixin, CreateView):
    model = IncidentType
    template_name = 'incidenttype/incidentcreate.html'
    fields = '__all__'
    success_url = reverse_lazy('incidenttype_list')
    allowed_roles = ['National', 'Regional']


class IncidentTypeDetailView(LoginRequiredMixin, RoleRequiredMixin, DetailView):
    model = IncidentType
    template_name = 'incidenttype/incidentdetail.html'
    context_object_name = 'type'
    allowed_roles = ['National', 'Regional']


class IncidentTypeUpdateView(LoginRequiredMixin, RoleRequiredMixin, UpdateView):
    model = IncidentType
    template_name = 'incidenttype/incidentupdate.html'
    fields = '__all__'
    success_url = reverse_lazy('incidenttype_list')
    allowed_roles = ['National', 'Regional']


class IncidentTypeDeleteView(LoginRequiredMixin, RoleRequiredMixin, DeleteView):
    model = IncidentType
    template_name = 'incidenttype/incidentdelete.html'
    success_url = reverse_lazy('incidenttype_list')
    allowed_roles = ['National', 'Regional']


class SanitaryIncidentListView(LoginRequiredMixin, RoleRequiredMixin, ListView):
    model = SanitaryIncident
    template_name = 'sanitaryincident/list.html'
    context_object_name = 'incidents'
    paginate_by = 10
    ordering = ['-date_time']
    allowed_roles = ['National', 'Regional']

    def get_queryset(self):
        queryset = SanitaryIncident.objects.filter(status='validated').order_by(*self.ordering)
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(description__icontains=search) |
                Q(city__name__icontains=search))

        if self.request.GET.get('incident_type'):
            queryset = queryset.filter(incident_type_id=self.request.GET.get('incident_type'))

        if self.request.GET.get('outcome'):
            queryset = queryset.filter(outcome=self.request.GET.get('outcome'))

        if self.request.GET.get('date'):
            queryset = queryset.filter(date_time__date=self.request.GET.get('date'))

        if self.request.GET.get('city'):
            queryset = queryset.filter(city_id=self.request.GET.get('city'))

        if self.request.GET.get('event'):
            queryset = queryset.filter(event_id=self.request.GET.get('event'))

        min_people = self.request.GET.get('min_people')
        max_people = self.request.GET.get('max_people')
        if min_people:
            queryset = queryset.filter(number_of_people_involved__gte=min_people)
        if max_people:
            queryset = queryset.filter(number_of_people_involved__lte=max_people)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['incident_types'] = IncidentType.objects.all()
        context['cities'] = Commune.objects.all()
        context['events'] = MajorEvent.objects.all()
        context['outcome_choices'] = SanitaryIncident.OUTCOME_CHOICES
        return context


class IncidentToValidListView(LoginRequiredMixin, RoleRequiredMixin, ListView):
    model = SanitaryIncident
    template_name = 'sanitaryincident/non_valid_list.html'
    context_object_name = 'incidents'
    paginate_by = 10
    ordering = ['-date_time']
    allowed_roles = ['National', 'Regional']

    def get_queryset(self):
        return SanitaryIncident.objects.exclude(status='validated').order_by(*self.ordering)


class SanitaryIncidentCreateView(LoginRequiredMixin, RoleRequiredMixin, CreateView):
    model = SanitaryIncident
    template_name = 'sanitaryincident/create.html'
    form_class = SanitaryIncidentForm
    success_url = reverse_lazy('sanitaryincident_non_valid_list')
    allowed_roles = ['National', 'Regional']

    def form_valid(self, form):
        # 1. Sauvegarde sans commit pour ajouter l'événement
        incident = form.save(commit=False)

        # 2. Recherche d’un événement en cours à la date de l’incident
        incident_date = form.cleaned_data.get('date_time')

        matching_event = MajorEvent.objects.filter(
            start_date__lte=incident_date,
            end_date__gte=incident_date
        ).first()

        if matching_event:
            incident.event = matching_event

        # 3. Ajout de l'utilisateur qui poste si besoin
        incident.posted_by = self.request.user if self.request.user.is_authenticated else None

        # 4. Sauvegarde finale
        incident.save()
        form.save_m2m()

        messages.success(self.request, 'Incident enregistré avec succès !')
        return redirect(self.success_url)

    def form_invalid(self, form):
        messages.error(self.request, 'Veuillez corriger les erreurs ci-dessous :')

        # Boucle sur les champs pour afficher chaque erreur individuellement
        for field, errors in form.errors.items():
            field_label = form.fields.get(field).label if field in form.fields else field
            for error in errors:
                messages.error(self.request, format_html("<strong>{}</strong>: {}", field_label, error))

        return super().form_invalid(form)


@require_POST
def validate_incident(request, pk):
    incident = get_object_or_404(SanitaryIncident, pk=pk)
    incident.status = 'validated'
    incident.save()
    messages.success(request, "✅ Incident validé avec succès.")
    return redirect('sanitaryincident_detail', pk=pk)


@require_POST
def reject_incident(request, pk):
    incident = get_object_or_404(SanitaryIncident, pk=pk)
    incident.status = 'rejected'
    incident.save()
    messages.warning(request, "🚫 Incident rejeté.")
    return redirect('sanitaryincident_detail', pk=pk)


# class IncidentMapView(RoleRequiredMixin, TemplateView):
#     template_name = 'sanitaryincident/incident_map.html'
#     allowed_roles = ['National', 'Regional']
#     redirect_view_if_denied = 'public_dashboard'
#
#     def get_context_data(self, **kwargs):
#         context = super().get_context_data(**kwargs)
#         incidents = SanitaryIncident.objects.all()
#         context['incidents_geojson'] = serialize('geojson', incidents,
#                                                  geometry_field='location',
#                                                  fields=(
#                                                      'id', 'incident_type__name', 'status', 'date_time', 'city__name'))
#         return context
#
#     def get(self, request, *args, **kwargs):
#         if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
#             if request.GET.get("layer") == "districts":
#                 return JsonResponse(self.get_districts_geojson())
#
#         if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
#             incidents = SanitaryIncident.objects.all()
#             data = {
#                 'type': 'FeatureCollection',
#                 'features': [{
#                     'type': 'Feature',
#                     'geometry': {
#                         'type': 'Point',
#                         'coordinates': [incident.location.x, incident.location.y] if incident.location else None
#                     },
#                     'properties': {
#                         'id': incident.id,
#                         'type': incident.incident_type.name,
#                         'status': incident.get_status_display(),
#                         'date': incident.date_time.strftime('%d/%m/%Y %H:%M'),
#                         'location': incident.city.name if incident.city else 'Inconnu',
#                         'outcome': incident.get_outcome_display(),
#                         'people_involved': incident.number_of_people_involved,
#                         'icon': self.get_incident_icon(incident)
#                     }
#                 } for incident in incidents if incident.location]
#             }
#             return JsonResponse(data)
#         return super().get(request, *args, **kwargs)
#
#     def get_incident_icon(self, incident):
#         if incident.status == 'validated':
#             return 'validated-icon'
#         elif incident.status == 'rejected':
#             return 'rejected-icon'
#         return 'pending-icon'
#
#     def get_districts_geojson(self):
#         from cogu.models import DistrictSanitaire
#
#         qs = DistrictSanitaire.objects.annotate(
#             incident_count=Count('commune__sanitaryincident')
#         ).filter(geom__isnull=False)
#
#         features = []
#         for district in qs:
#             if district.geom:
#                 features.append({
#                     "type": "Feature",
#                     "geometry": json.loads(district.geom.geojson),
#                     "properties": {
#                         "id": district.id,
#                         "name": district.nom,
#                         "region": district.region.name if district.region else None,
#                         "incident_count": district.incident_count
#                     }
#                 })
#
#         return {
#             "type": "FeatureCollection",
#             "features": features
#         }

def export_incidents_csv(request):
    incidents = SanitaryIncident.objects.all()  # à filtrer selon les paramètres si nécessaire

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="incidents.csv"'

    writer = csv.writer(response)
    writer.writerow(['ID', 'Type', 'Date', 'Commune', 'Statut', 'Issue', 'Personnes'])

    for i in incidents:
        writer.writerow([
            i.id,
            i.incident_type.name,
            localtime(i.date_time).strftime("%d/%m/%Y %H:%M"),
            i.city.name if i.city else "Inconnu",
            i.get_status_display(),
            i.get_outcome_display(),
            i.number_of_people_involved
        ])

    return response


class IncidentMapView(RoleRequiredMixin, TemplateView):
    template_name = 'sanitaryincident/incident_map.html'
    allowed_roles = ['National', 'Regional']
    redirect_view_if_denied = 'public_dashboard'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        incidents = SanitaryIncident.objects.all()

        # Ajout des districts sanitaires au contexte
        # districts = DistrictSanitaire.objects.all()
        districts = DistrictSanitaire.objects.exclude(geom__isnull=True)
        districts = [d for d in districts if isinstance(d.geom, GEOSGeometry)]
        context['districts_geojson'] = serialize('geojson', districts,
                                                 geometry_field='geom',
                                                 fields=('id', 'nom', 'region__name'))

        context['incidents_geojson'] = serialize('geojson', incidents,
                                                 geometry_field='location',
                                                 fields=(
                                                     'id', 'incident_type__name', 'status', 'date_time', 'city__name'))
        # context['poles'] = PolesRegionaux.objects.all()
        # context['regions'] = HealthRegion.objects.select_related('poles')
        # context['districts'] = DistrictSanitaire.objects.select_related('region')
        # context['incident_types'] = IncidentType.objects.only('id', 'name')

        context['poles'] = PolesRegionaux.objects.only('id', 'name')
        context['regions'] = HealthRegion.objects.select_related('poles').only('id', 'name', 'poles__name')
        context['districts'] = DistrictSanitaire.objects.select_related('region').only('id', 'nom', 'region__name')
        context['incident_types'] = IncidentType.objects.only('id', 'name')

        return context

    def get(self, request, *args, **kwargs):
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            incidents = SanitaryIncident.objects.select_related(
                'city__district__region__poles', 'incident_type'
            ).all()

            # Filtres GET
            start_date = request.GET.get('start_date')
            end_date = request.GET.get('end_date')
            pole_id = request.GET.get('pole_id')
            region_id = request.GET.get('region_id')
            district_id = request.GET.get('district_id')
            incident_type__id = request.GET.get('type_id')

            if start_date:
                incidents = incidents.filter(date_time__date__gte=start_date)
            if end_date:
                incidents = incidents.filter(date_time__date__lte=end_date)
            if pole_id:
                incidents = incidents.filter(city__district__region__poles__id=pole_id)
            if region_id:
                incidents = incidents.filter(city__district__region__id=region_id)
            if district_id:
                incidents = incidents.filter(city__district__id=district_id)
            if incident_type__id:
                incidents = incidents.filter(incident_type=incident_type__id)

            # Regroupement des districts avec incidents filtrés
            districts = DistrictSanitaire.objects.exclude(geojson__isnull=True)
            district_incident_counts = {
                d.id: incidents.filter(city__district=d).count() for d in districts
            }

            # Construction du GeoJSON
            data = {
                'incidents': {
                    'type': 'FeatureCollection',
                    'features': [
                        {
                            'type': 'Feature',
                            'geometry': {
                                'type': 'Point',
                                'coordinates': [i.location.x, i.location.y] if i.location else None
                            },
                            'properties': {
                                'id': i.id,
                                'type': i.incident_type.name,
                                'status': i.get_status_display(),
                                'date': i.date_time.strftime('%d/%m/%Y %H:%M'),
                                'location': i.city.name if i.city else 'Inconnu',
                                'outcome': i.get_outcome_display(),
                                'people_involved': i.number_of_people_involved,
                                'icon': self.get_incident_icon(i),
                                'district_id': i.city.district.id if i.city and i.city.district else None
                            }
                        }
                        for i in incidents if i.location
                    ]
                },
                'districts': {
                    'type': 'FeatureCollection',
                    'features': [
                        {
                            'type': 'Feature',
                            'geometry': d.geojson.get("geometry"),
                            'properties': {
                                'id': d.id,
                                'name': d.nom,
                                'region': d.region.name if d.region else '',
                                'pole': d.region.poles.name if d.region and d.region.poles else '',
                                'incident_count': district_incident_counts.get(d.id, 0)
                            }
                        }
                        for d in districts if d.geojson and d.geojson.get("geometry")
                    ]
                }
            }

            return JsonResponse(data)
        return super().get(request, *args, **kwargs)

    def get_incident_icon(self, incident):
        if incident.status == 'validated':
            return 'validated-icon'
        elif incident.status == 'rejected':
            return 'rejected-icon'
        return 'pending-icon'


#     def get(self, request, *args, **kwargs):
#         if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
#             incidents = SanitaryIncident.objects.all()
#             districts = DistrictSanitaire.objects.exclude(geojson__isnull=True)
#
#             # Compter les incidents par district
#             district_incident_counts = {
#                 d.id: incidents.filter(city__district=d).count() for d in districts
#             }
#
#             data = {
#                 'incidents': {
#                     'type': 'FeatureCollection',
#                     'features': [
#                         {
#                             'type': 'Feature',
#                             'geometry': {
#                                 'type': 'Point',
#                                 'coordinates': [i.location.x, i.location.y] if i.location else None
#                             },
#                             'properties': {
#                                 'id': i.id,
#                                 'type': i.incident_type.name,
#                                 'status': i.get_status_display(),
#                                 'date': i.date_time.strftime('%d/%m/%Y %H:%M'),
#                                 'location': i.city.name if i.city else 'Inconnu',
#                                 'outcome': i.get_outcome_display(),
#                                 'people_involved': i.number_of_people_involved,
#                                 'icon': self.get_incident_icon(i),
#                                 'district_id': i.city.district.id if i.city and i.city.district else None
#                             }
#                         }
#                         for i in incidents if i.location
#                     ]
#                 },
#                 'districts': {
#                     'type': 'FeatureCollection',
#                     'features': [
#                         {
#                             'type': 'Feature',
#                             'geometry': d.geojson['geometry'],  # ✅ ici on extrait uniquement la géométrie
#                             'properties': {
#                                 'id': d.id,
#                                 'name': d.nom,
#                                 'region': d.region.name if d.region else '',
#                                 'incident_count': district_incident_counts.get(d.id, 0)
#                             }
#                         }
#                         for d in districts
#         if d.geojson and d.geojson.get("geometry")  # ✅ s'assurer que c'est un Feature
#     ]
# }
#             }
#
#             return JsonResponse(data)
#         return super().get(request, *args, **kwargs)


class SanitaryIncidentDetailView(LoginRequiredMixin, RoleRequiredMixin, DetailView):
    model = SanitaryIncident
    template_name = 'sanitaryincident/detail.html'
    context_object_name = 'incident'
    allowed_roles = ['National', 'Regional']
    redirect_view_if_denied = 'public_dashboard'

    def get_queryset(self):
        return (
            SanitaryIncident.objects
            .select_related('message', 'city', 'incident_type')
            .prefetch_related('media', 'patients_related')
            # .filter(status='validated')  # filtre si on veut afficher seulement les incidents validés
        )


class SanitaryIncidentUpdateView(LoginRequiredMixin, RoleRequiredMixin, UpdateView):
    model = SanitaryIncident
    template_name = 'sanitaryincident/update.html'
    form_class = SanitaryIncidentForm
    success_url = reverse_lazy('sanitaryincident_list')
    allowed_roles = ['National', 'Regional']
    redirect_view_if_denied = 'public_dashboard'

    def form_valid(self, form):
        messages.success(self.request, 'Incident enregistré avec succès!')
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, 'Veuillez corriger les erreurs ci-dessous :')

        # Boucle sur les champs pour afficher chaque erreur individuellement
        for field, errors in form.errors.items():
            field_label = form.fields.get(field).label if field in form.fields else field
            for error in errors:
                messages.error(self.request, format_html("<strong>{}</strong>: {}", field_label, error))

        return super().form_invalid(form)


class SanitaryIncidentDeleteView(LoginRequiredMixin, RoleRequiredMixin, DeleteView):
    model = SanitaryIncident
    template_name = 'sanitaryincident/delete.html'
    success_url = reverse_lazy('sanitaryincident_list')
    allowed_roles = ['National', 'Regional']
    redirect_view_if_denied = 'public_dashboard'


class WhatsAppMessageListView(LoginRequiredMixin, RoleRequiredMixin, ListView):
    model = WhatsAppMessage
    template_name = 'pages/whatsapp/messages_list.html'
    context_object_name = 'messages'
    paginate_by = 20
    ordering = ['-timestamp']
    allowed_roles = ['National', 'Regional']
    redirect_view_if_denied = 'public_dashboard'

    def get_queryset(self):
        return WhatsAppMessage.objects.all().order_by(*self.ordering)


def contact(request):
    if request.method == 'POST':
        form = ContactForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Votre message a été envoyé avec succès. Nous vous contacterons bientôt!")
            return redirect('contact')  # Replace with your actual contact URL name
        else:
            messages.error(request, "Il y a eu une erreur dans l'envoi de votre message. Veuillez réessayer.")
    else:
        form = ContactForm()

    return render(request, 'pages/landing.html', {'form': form})


# def generate_cogu_report(request, *args, **kwargs):
#     # today = timezone.now().date()
#     today = timezone.now().date()
#     yesterday = today - timedelta(days=1)
#     today = yesterday
#
#     output_format = request.GET.get('format') or kwargs.get('format', 'pdf')
#
#     daily_incidents = SanitaryIncident.objects.filter(
#         date_time__date=today
#     ).select_related(
#         'incident_type', 'city__district__region'
#     )
#
#     total_incidents = daily_incidents.count()
#     validated_incidents = daily_incidents.filter(status='validated').count()
#     pending_incidents = daily_incidents.filter(status='pending').count()
#
#     resolved_incidents = SanitaryIncident.objects.filter(
#         date_time__date=yesterday,
#         status='validated'
#     ).count()
#
#     regions = HealthRegion.objects.all()
#     region_data = []
#
#     for region in regions:
#         region_incidents = daily_incidents.filter(city__district__region=region)
#
#         incident_types = {}
#         for incident in region_incidents:
#             type_name = incident.incident_type.name
#             if type_name not in incident_types:
#                 incident_types[type_name] = {
#                     'validated': 0,
#                     'pending': 0,
#                     'total': 0
#                 }
#             incident_types[type_name]['total'] += 1
#             if incident.status == 'validated':
#                 incident_types[type_name]['validated'] += 1
#             else:
#                 incident_types[type_name]['pending'] += 1
#
#         region_data.append({
#             'name': region.name,
#             'total_incidents': region_incidents.count(),
#             'incident_types': incident_types,
#             'actions': get_actions_for_region(region.name)
#         })
#
#     context = {
#         'date': today.strftime("%d %B %Y"),
#         'total_incidents': total_incidents,
#         'validated_incidents': validated_incidents,
#         'pending_incidents': pending_incidents,
#         'resolved_incidents': resolved_incidents,
#         'region_data': region_data,
#         'actions_taken': get_actions_taken(),
#         'recommendations': get_recommendations(),
#         'next_steps': get_next_steps(),
#         'logo_armoirie_path': '/static/assets/media/armoirie_ci.png',
#         'logo_sante_path': '/static/assets/media/logoMSHPCMU.png',
#         'logo_afriqconsulting_path': '/static/assets/media/logo-AFRIQ-CONSULTING.png',
#     }
#
#     if output_format == 'pdf':
#         return generate_pdf_report(context)
#     elif output_format == 'word':
#         return generate_word_report(context)
#     else:
#         return HttpResponse("Invalid format specified", status=400)
#
def generate_cogu_report(request):
    if request.method == 'POST':
        form = CoguReportForm(request.POST)
        if form.is_valid():
            start_date = form.cleaned_data['start_date']
            end_date = form.cleaned_data['end_date']
            comments = form.cleaned_data['comments']
            output_format = form.cleaned_data['format']

            context = prepare_report_context(start_date, end_date, comments)

            if output_format == 'pdf':
                result = generate_pdf_report(context, return_bytes=True)
                if result:
                    report = CoguReport.objects.create(
                        created_by=request.user,
                        report_date=timezone.now().date(),
                        start_date=start_date,
                        end_date=end_date,
                        comments=comments,
                        format=output_format
                    )
                    filename = f"COGU_Report_{start_date.strftime('%Y%m%d')}_{end_date.strftime('%Y%m%d')}.pdf"
                    report.file.save(filename, ContentFile(result))
                    return HttpResponse(result, content_type='application/pdf')
                return HttpResponse('Erreur génération PDF', status=500)

            elif output_format == 'word':
                response = generate_word_report(context)
                return response

            else:
                return HttpResponse("Invalid format specified", status=400)
    else:
        form = CoguReportForm()

    return render(request, 'reports/generate_report_form.html', {'form': form})


def prepare_report_context(start_date, end_date, comments=None):
    # Filtre de base pour les incidents
    incidents = SanitaryIncident.objects.filter(
        date_time__date__range=[start_date, end_date]
    ).select_related(
        'incident_type', 'city__district__region', 'event'
    )

    # Statistiques globales
    stats = incidents.aggregate(
        total_incidents=Count('id'),
        total_victims=Sum('number_of_people_involved'),
        total_deaths=Sum('deces_nbr'),
        total_injuries=Sum('blessure_nbr'),
        total_evacuations=Sum('evacues_nbr'),
        total_treated=Sum('pris_en_charge_nbr'),
        total_exeat=Sum('exeat_nbr'),
        validated_incidents=Count('id', filter=models.Q(status='validated')),
        pending_incidents=Count('id', filter=models.Q(status='pending'))
    )

    # Statistiques par événement majeur
    events_stats = []
    major_events = MajorEvent.objects.filter(
        start_date__lte=end_date,
        end_date__gte=start_date
    )

    for event in major_events:
        event_incidents = incidents.filter(event=event)
        event_stats = event_incidents.aggregate(
            count=Count('id'),
            deaths=Sum('deces_nbr'),
            injuries=Sum('blessure_nbr'),
            evacuations=Sum('evacues_nbr')
        )
        events_stats.append({
            'event': event,
            'stats': event_stats
        })

    # Statistiques par pôle régional
    poles_stats = []
    for pole in PolesRegionaux.objects.all():
        pole_incidents = incidents.filter(city__district__region__poles=pole)
        pole_stats = pole_incidents.aggregate(
            count=Count('id'),
            deaths=Sum('deces_nbr'),
            injuries=Sum('blessure_nbr')
        )
        poles_stats.append({
            'pole': pole,
            'stats': pole_stats
        })

    # Statistiques par région et district
    regions_data = []
    districts_geojson = {
        "type": "FeatureCollection",
        "features": []
    }

    for region in HealthRegion.objects.all():
        region_incidents = incidents.filter(city__district__region=region)

        # Stats par type d'incident
        incident_types = defaultdict(lambda: {
            'total': 0, 'validated': 0, 'pending': 0,
            'deaths': 0, 'injuries': 0, 'evacuations': 0
        })

        for incident in region_incidents:
            type_name = incident.incident_type.name
            incident_types[type_name]['total'] += 1
            if incident.status == 'validated':
                incident_types[type_name]['validated'] += 1
            else:
                incident_types[type_name]['pending'] += 1

            incident_types[type_name]['deaths'] += incident.deces_nbr
            incident_types[type_name]['injuries'] += incident.blessure_nbr
            incident_types[type_name]['evacuations'] += incident.evacues_nbr

        # Stats par district
        districts_data = []
        for district in region.districts.all():
            district_incidents = region_incidents.filter(city__district=district)
            district_stats = district_incidents.aggregate(
                count=Count('id'),
                deaths=Sum('deces_nbr'),
                injuries=Sum('blessure_nbr'),
                evacuations=Sum('evacues_nbr')
            )

            districts_data.append({
                'district': district,
                'stats': district_stats
            })

            # Ajout au GeoJSON si le district a des incidents
            if district.geojson and district_stats['count'] > 0:
                districts_geojson['features'].append({
                    "type": "Feature",
                    "geometry": district.geojson.get("geometry"),
                    "properties": {
                        "name": district.nom,
                        "incidents": district_stats['count'],
                        "deaths": district_stats['deaths'] or 0,
                        "injuries": district_stats['injuries'] or 0
                    }
                })

        regions_data.append({
            'region': region,
            'stats': {
                'total': region_incidents.count(),
                'deaths': region_incidents.aggregate(s=Sum('deces_nbr'))['s'] or 0,
                'injuries': region_incidents.aggregate(s=Sum('blessure_nbr'))['s'] or 0
            },
            'incident_types': dict(incident_types),
            'districts': districts_data,
            'actions': get_actions_for_region(region.name)
        })

    # Préparation des données pour les graphiques
    chart_data = {'incidents_by_type': list(
        incidents.values('incident_type__name').annotate(count=Count('id')).order_by('-count')[:5]),
                  'incidents_by_region': list(
                      incidents.values('city__district__region__name').annotate(count=Count('id')))}
    districts_geojson_str = json.dumps(districts_geojson)
    map_image = generate_map_image(
        json.loads(districts_geojson_str))  # ou directement : generate_map_image(districts_geojson)
    return {
        'start_date': start_date.strftime("%d %B %Y"),
        'end_date': end_date.strftime("%d %B %Y"),
        'report_date': timezone.now().date().strftime("%d %B %Y"),
        'global_stats': stats,
        'events_stats': events_stats,
        'poles_stats': poles_stats,
        'regions_data': regions_data,
        'districts_geojson': json.dumps(districts_geojson),
        'map_image':   map_image,
        'chart_data': chart_data,
        'actions_taken': get_actions_taken(),
        'recommendations': get_recommendations(),
        'next_steps': get_next_steps(),
        'comments': comments,
        'logo_paths': {
            'armoirie': '/static/assets/media/armoirie_ci.png',
            'sante': '/static/assets/media/logoMSHPCMU.png',
            'afriqconsulting': '/static/assets/media/logo-AFRIQ-CONSULTING.png'
        }
    }


def generate_pdf_report(context, return_bytes=False):
    context['chart_type_image'] = generate_chart_image(
        data=context['chart_data']['incidents_by_type'],
        title="Répartition par type d'incident",
        x_field='incident_type__name',
        y_field='count'
    )

    context['chart_region_image'] = generate_chart_image(
        data=context['chart_data']['incidents_by_region'],
        title="Incidents par région",
        x_field='city__district__region__name',
        y_field='count'
    )

    context['map_image'] = generate_map_image(json.loads(context['districts_geojson']))

    template = get_template('reports/cogu_report.html')
    html = template.render(context)
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)

    if pdf.err:
        return None

    if return_bytes:
        return result.getvalue()

    response = HttpResponse(result.getvalue(), content_type='application/pdf')
    filename = f"COGU_Report_{context['start_date']}_{context['end_date']}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def generate_chart_image(data, title, x_field, y_field):
    if not data:
        return ""  # ou une image par défaut

    try:
        plt.figure(figsize=(8, 4))
        x = [item[x_field] for item in data]
        y = [item[y_field] for item in data]
        plt.bar(x, y)
        plt.title(title)
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150)
        plt.close()
        buf.seek(0)

        return base64.b64encode(buf.read()).decode('utf-8')
    except Exception as e:
        print(f"Erreur lors de la génération du graphique : {e}")
        return ""


# def generate_map_image(geojson_data):
#     """Génère une image de carte à partir des données GeoJSON"""
#     # Utilisez la même fonction que précédemment avec selenium
#     # ou une alternative comme folium pour générer une image
#     temp_path = os.path.join(settings.MEDIA_ROOT, 'temp_map.png')
#     generate_map_image_with_selenium(geojson_data, temp_path)
#
#     with open(temp_path, 'rb') as f:
#         img_data = base64.b64encode(f.read()).decode('utf-8')
#
#     os.remove(temp_path)
#     return img_data
def generate_map_image(geojson_data):
    try:
        def generate_map_image_with_selenium(geojson_data, output_path):
            import folium
            from selenium import webdriver
            from selenium.webdriver.chrome.options import Options
            import time

            # Création de la carte
            m = folium.Map(location=[7.54, -5.55], zoom_start=6)

            # Ajout des polygones avec couleurs dynamiques
            def style_function(feature):
                incidents = feature['properties'].get('incidents', 0)
                color = (
                    '#ff0000' if incidents > 10 else
                    '#ffa500' if incidents > 5 else
                    '#ffff00' if incidents > 0 else
                    '#d3d3d3'
                )
                return {
                    'fillColor': color,
                    'color': 'black',
                    'weight': 1,
                    'fillOpacity': 0.6,
                }

            folium.GeoJson(
                geojson_data,
                name="Districts",
                style_function=style_function,
                tooltip=folium.GeoJsonTooltip(fields=['name', 'incidents', 'deaths', 'injuries'])
            ).add_to(m)

            # Sauvegarde HTML temporaire
            html_path = os.path.join(settings.MEDIA_ROOT, 'temp_map.html')
            m.save(html_path)

            # Capture avec Selenium
            options = Options()
            options.add_argument('--headless')
            options.add_argument('--disable-gpu')
            options.add_argument('--window-size=1200x800')
            driver = webdriver.Chrome(options=options)
            driver.get(f'file://{html_path}')
            time.sleep(2)
            driver.save_screenshot(output_path)
            driver.quit()
            os.remove(html_path)
        # def generate_map_image_with_selenium(geojson_data, output_path):
        #     m = folium.Map(location=[7.54, -5.55], zoom_start=6)
        #     folium.GeoJson(
        #         geojson_data,
        #         name="Districts",
        #         style_function=lambda feature: {
        #             'fillColor': '#ff0000' if feature['properties']['incidents'] > 10 else (
        #                 '#ffa500' if feature['properties']['incidents'] > 5 else (
        #                     '#ffff00' if feature['properties']['incidents'] > 0 else '#d3d3d3'
        #                 )
        #             ),
        #             'color': 'black',
        #             'weight': 1,
        #             'fillOpacity': 0.6,
        #         },
        #         tooltip=folium.GeoJsonTooltip(fields=['name', 'incidents', 'deaths', 'injuries'])
        #     ).add_to(m)
        #     html_path = os.path.join(settings.MEDIA_ROOT, 'temp_map.html')
        #     m.save(html_path)
        #
        #     options = Options()
        #     options.add_argument('--headless')
        #     options.add_argument('--disable-gpu')
        #     options.add_argument('--window-size=1200x800')
        #     driver = webdriver.Chrome(options=options)
        #     driver.get(f'file://{html_path}')
        #     time.sleep(2)
        #     driver.save_screenshot(output_path)
        #     driver.quit()
        #     os.remove(html_path)

        output_path = os.path.join(settings.MEDIA_ROOT, 'temp_map.png')
        generate_map_image_with_selenium(geojson_data, output_path)

        with open(output_path, 'rb') as f:
            img_data = base64.b64encode(f.read()).decode('utf-8')

        os.remove(output_path)
        return img_data

    except Exception as e:
        print(f"Erreur lors de la génération de la carte : {e}")
        return ""


def generate_word_report(context):
    document = Document()

    # ==================== EN-TÊTE AVEC LOGOS ====================
    section = document.sections[0]
    header = section.header

    # Créer un tableau pour organiser les logos et le titre
    table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    table.autofit = False

    # Configurer les largeurs des colonnes
    widths = (Inches(1.5), Inches(3.5), Inches(1.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Cellule gauche - Logo des armoiries
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    try:
        left_para.add_run().add_picture(
            os.path.join(settings.STATIC_ROOT, 'assets/media/images.jpeg'),
            width=Inches(1.3)
        )
    except:
        left_para.add_run("[LOGO ARMOIRIES CIV]")

    # Cellule centrale - Titre
    center_cell = table.cell(0, 1)
    center_para = center_cell.paragraphs[0]
    center_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title_run = center_para.add_run('RAPPORT JOURNALIER COGU')
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = center_para.add_run('\nMinistère de la Santé et de l\'Hygiène Publique')
    subtitle.font.size = Pt(12)

    # Cellule droite - Logo ministère santé
    right_cell = table.cell(0, 2)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    try:
        right_para.add_run().add_picture(
            os.path.join(settings.STATIC_ROOT, 'assets/media/logoMSHPCMU.png'),
            width=Inches(1.3)
        )
    except:
        right_para.add_run("[LOGO MINISTERE SANTE]")

    # Ajouter une ligne de séparation
    header.add_paragraph()
    p = header.add_paragraph()
    p.add_run().add_break()
    p.add_run("_" * 100).bold = True

    # ==================== CORPS DU DOCUMENT ====================
    document.add_paragraph(f"Date : {context['date']}")
    document.add_paragraph("Destinataires : Monsieur le Ministre de la Santé, Membres du COGU")
    document.add_paragraph("Émetteur : Directeur Général de la Santé et de l'Hygiène Publique (DGSHP)")

    # document.add_heading('RAPPORT JOURNALIER COGU', level=0)
    # document.add_paragraph(f"Date : {context['date']}")
    # document.add_paragraph("Destinataires : Monsieur le Ministre de la Santé, Membres du COGU")
    # document.add_paragraph("Émetteur : Directeur Général de la Santé et de l'Hygiène Publique (DGSHP)")

    document.add_heading('RÉCAPITULATIF GLOBAL', level=1)
    document.add_paragraph(f"Nombre total d'incidents signalés : {context['total_incidents']}")
    document.add_paragraph(f"Incidents validés : {context['validated_incidents']}")
    document.add_paragraph(f"Incidents en cours de validation : {context['pending_incidents']}")
    document.add_paragraph(f"Incidents résolus hier : {context['resolved_incidents']}")

    document.add_heading('DÉTAILS PAR RÉGION SANITAIRE', level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Région'
    hdr_cells[1].text = 'Total'
    hdr_cells[2].text = 'Types d’incidents'
    hdr_cells[3].text = 'Statuts'
    hdr_cells[4].text = 'Actions'

    for region in context['region_data']:
        row = table.add_row().cells
        row[0].text = region['name']
        row[1].text = str(region['total_incidents'])
        types = [f"- {c['total']} cas de {t}" for t, c in region['incident_types'].items()]
        row[2].text = "\n".join(types)
        statuts = [f"- {c['validated']} validés, {c['pending']} en cours" for t, c in region['incident_types'].items()]
        row[3].text = "\n".join(statuts)
        row[4].text = "\n".join([f"- {a}" for a in region['actions']])

    for title, items in [
        ("ACTIONS MENÉES ET INTERVENTIONS EN COURS", context['actions_taken']),
        ("RECOMMANDATIONS ET PERSPECTIVES", context['recommendations']),
        ("PROCHAINES ÉTAPES", context['next_steps']),
    ]:
        document.add_heading(title, level=1)
        for item in items:
            document.add_paragraph(item, style='List Bullet')

    document.add_heading('Conclusion', level=1)
    document.add_paragraph(
        "La situation reste sous contrôle, avec un bon niveau de réactivité des équipes de terrain. "
        "Les investigations se poursuivent. Un suivi quotidien est maintenu pour informer Monsieur le Ministre."
    )
    document.add_paragraph(f"Fait à Abidjan, le {context['date']}")
    document.add_paragraph("Signature :")
    document.add_paragraph("Pr. SAMBA Mamadou")
    document.add_paragraph("Directeur Général de la Santé et de l'Hygiène Publique")

    # ==================== PIED DE PAGE ====================
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_para.text = f"© {datetime.now().year} Ministère de la Santé - Tous droits réservés"

    # Ajouter le logo Afriq Consulting
    footer_para.add_run().add_break()
    try:
        footer_para.add_run().add_picture(
            os.path.join(settings.STATIC_ROOT, 'assets/media/logo-AFRIQ-CONSULTING.png'),
            width=Inches(1.0)
        )
        footer_para.add_run(" - Solution développée par Afriq Consulting")
    except:
        footer_para.add_run("[LOGO AFRIQ CONSULTING]")

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    response = HttpResponse(file_stream.read(),
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="COGU_Report_{context["date"]}.docx"'
    return response


# Helper functions
def get_actions_for_region(region_name):
    # This would be customized based on your business logic
    actions = {
        'Abidjan 1': [
            "Enquête épidémiologique",
            "Prélèvements et analyses en cours"
        ],
        'Gbêkê (Bouaké)': [
            "Équipe de surveillance déployée",
            "Sensibilisation communautaire"
        ],
        'Haut-Sassandra': [
            "Traitement symptomatique",
            "Aucune complication signalée"
        ],
        'Poro (Korhogo)': [
            "Mise en observation des cas",
            "Campagne d'hygiène alimentaire"
        ],
        'Sud-Comoé (Aboisso)': [
            "Investigation médicale",
            "Surveillance renforcée"
        ]
    }
    return actions.get(region_name, ["Actions en cours d'évaluation"])


def get_actions_taken():
    return [
        "Des équipes multidisciplinaires (épidémiologistes, agents communautaires) sont actuellement sur le terrain pour confirmer ou infirmer les cas de maladies à potentiel épidémique.",
        "Les districts concernés ont reçu des directives pour intensifier la surveillance épidémiologique dans les formations sanitaires voisines.",
        "Une mise à jour du protocole d'investigation a été transmise à tous les agents de santé.",
        "Campagnes de sensibilisation en cours dans les zones touchées (mesures d'hygiène, importance du lavage des mains, consommation d'eau potable).",
        "Distribution de kits de chlore dans les localités à risque de choléra.",
        "Les équipes de coordination du COGU restent en contact permanent avec les responsables de districts.",
        "Des briefings quotidiens ont lieu pour actualiser la situation et définir les actions prioritaires."
    ]


def get_recommendations():
    return [
        "Intensifier la communication locale et les activités de promotion de la santé (radio, SMS, affiches) pour prévenir la propagation de maladies infectieuses.",
        "Réduire le délai de validation des alertes en attente afin de permettre une réaction rapide et d'éviter tout retard dans la prise en charge.",
        "Impliquer davantage les collectivités locales (chefs de village, leaders communautaires) pour identifier rapidement les nouveaux cas et encourager la vaccination (le cas échéant).",
        "Vérifier la disponibilité des stocks de médicaments essentiels et de matériel médical dans les centres de santé concernés."
    ]


def get_next_steps():
    return [
        "Finalisation des analyses : Les laboratoires régionaux enverront leurs résultats dans un délai de 48 heures pour confirmer le diagnostic des cas suspects.",
        "Renforcement de la vaccination : En cas de confirmation de maladies épidémiques, une campagne de vaccination ou de traitement préventif sera envisagée en priorité dans les zones touchées.",
        "Prochain rapport : Un nouveau point de situation sera diffusé demain à la même heure pour tous les membres du COGU et le cabinet du Ministre."
    ]


class CoguReportListView(ListView):
    model = CoguReport
    template_name = 'reports/report_list.html'
    context_object_name = 'reports'
    paginate_by = 10
    ordering = ['-created_at']


def download_report(request, pk):
    report = get_object_or_404(CoguReport, pk=pk)
    # Ici, vous devrez implémenter la logique pour regénérer ou servir le rapport
    # Cela dépend de comment vous stockez les fichiers

    # Exemple simplifié:
    context = prepare_report_context(report.start_date, report.end_date, report.comments)
    if report.format == 'pdf':
        return generate_pdf_report(context)
    else:
        return generate_word_report(context)


def view_report(request, pk):
    report = get_object_or_404(CoguReport, pk=pk)
    context = {
        'report': report,
        'details': prepare_report_context(report.start_date, report.end_date, report.comments)
    }
    return render(request, 'reports/report_details.html', context)


class IncidentReportView(LoginRequiredMixin, TemplateView):
    template_name = "reports/incident_report.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Récupération des paramètres de filtre
        date_from = self.request.GET.get('date_from')
        date_to = self.request.GET.get('date_to')
        status = self.request.GET.get('status')
        incident_type = self.request.GET.get('incident_type')
        region = self.request.GET.get('region')
        district = self.request.GET.get('district')
        severity = self.request.GET.get('severity')
        outcome = self.request.GET.get('outcome')

        # Filtrage de base
        incidents = SanitaryIncident.objects.all().select_related(
            'incident_type', 'city', 'city__district', 'city__district__region'
        ).prefetch_related('patients_related')

        # Application des filtres
        if date_from and date_to:
            try:
                date_from = datetime.strptime(date_from, '%Y-%m-%d')
                date_to = datetime.strptime(date_to, '%Y-%m-%d')
                incidents = incidents.filter(
                    date_time__date__range=[date_from, date_to]
                )
            except ValueError:
                pass
        else:
            # Par défaut, afficher les 30 derniers jours
            default_from = timezone.now() - timedelta(days=30)
            incidents = incidents.filter(date_time__gte=default_from)

        if status:
            incidents = incidents.filter(status=status)

        if incident_type:
            incidents = incidents.filter(incident_type_id=incident_type)

        if region:
            incidents = incidents.filter(city__district__region_id=region)

        if district:
            incidents = incidents.filter(city__district_id=district)

        if severity:
            if severity == 'high':
                incidents = incidents.filter(Q(number_of_people_involved__gte=5) | Q(outcome='mort'))
            elif severity == 'medium':
                incidents = incidents.filter(number_of_people_involved__gte=2, number_of_people_involved__lt=5)
            else:
                incidents = incidents.filter(number_of_people_involved=1)

        if outcome:
            incidents = incidents.filter(outcome=outcome)

        # Préparation des données pour les graphiques
        incidents_by_type = incidents.values(
            'incident_type__name'
        ).annotate(
            count=Count('id')
        ).order_by('-count')

        incidents_by_region = incidents.values(
            'city__district__region__name'
        ).annotate(
            count=Count('id')
        ).order_by('-count')

        incidents_by_status = incidents.values(
            'status'
        ).annotate(
            count=Count('id')
        ).order_by('status')

        # Récupération des options de filtre
        incident_types = IncidentType.objects.all()
        regions = HealthRegion.objects.all()
        districts = DistrictSanitaire.objects.all()
        if region:
            districts = districts.filter(region_id=region)

        total_count = incidents.count()
        validated_count = incidents.filter(status='validated').count()
        pending_count = incidents.filter(status='pending').count()
        rejected_count = incidents.filter(status='rejected').count()

        # Ajout au contexte
        context.update({
            'form': CoguReportForm(),
            'total_count': total_count,
            'validated_count': validated_count,
            'pending_count': pending_count,
            'rejected_count': rejected_count,

            'incidents': incidents.order_by('-date_time'),
            'incident_types': incident_types,
            'regions': regions,
            'districts': districts,
            'incidents_by_type': incidents_by_type,
            'incidents_by_region': incidents_by_region,
            'incidents_by_status': incidents_by_status,
            'filter_params': {
                'date_from': date_from.strftime('%Y-%m-%d') if date_from else '',
                'date_to': date_to.strftime('%Y-%m-%d') if date_to else '',
                'status': status,
                'incident_type': incident_type,
                'region': region,
                'district': district,
                'severity': severity,
                'outcome': outcome,
            },
            'status_choices': SanitaryIncident.STATUS_CHOICES,
            'outcome_choices': [
                ('mort', 'Décès'),
                ('blessure', 'Blessure'),
                ('sauvé', 'Sauvé'),
                ('autre', 'Autre'),
            ],
            'severity_choices': [
                ('high', 'Élevée'),
                ('medium', 'Moyenne'),
                ('low', 'Faible'),
            ],
        })

        return context


# Gestion des kits

class FournisseurListView(ListView):
    model = Fournisseur
    template_name = 'kits/fournisseur_list.html'


class FournisseurDetailView(DetailView):
    model = Fournisseur
    template_name = 'kits/fournisseur_detail.html'


class FournisseurCreateView(CreateView):
    model = Fournisseur
    fields = '__all__'
    template_name = 'kits/fournisseur_form.html'
    success_url = reverse_lazy('fournisseur_list')


class FournisseurUpdateView(UpdateView):
    model = Fournisseur
    fields = '__all__'
    template_name = 'kits/fournisseur_form.html'
    success_url = reverse_lazy('fournisseur_list')


class FournisseurDeleteView(DeleteView):
    model = Fournisseur
    template_name = 'kits/fournisseur_confirm_delete.html'
    success_url = reverse_lazy('fournisseur_list')


class KitListView(ListView):
    model = Kit
    template_name = 'kits/kit_list.html'


class KitDetailView(DetailView):
    model = Kit
    template_name = 'kits/kit_detail.html'


class KitCreateView(CreateView):
    model = Kit
    fields = '__all__'
    template_name = 'kits/kit_form.html'
    success_url = reverse_lazy('kit_list')


class KitUpdateView(UpdateView):
    model = Kit
    fields = '__all__'
    template_name = 'kits/kit_form.html'
    success_url = reverse_lazy('kit_list')


class KitDeleteView(DeleteView):
    model = Kit
    template_name = 'kits/kit_confirm_delete.html'
    success_url = reverse_lazy('kit_list')


# class StockDistrictView(TemplateView):
#     template_name = 'kits/stock_district.html'
#
#     def get_context_data(self, **kwargs):
#         context = super().get_context_data(**kwargs)
#
#         # Agrégation par district
#         stocks = (
#             Stock.objects
#             .select_related('centre__region__district')
#             .values(
#                 district_id=F('centre__region__district__id'),
#                 district_nom=F('centre__region__district__nom'),
#                 composant_nom=F('composant__nom'),
#                 unite=F('composant__unite_mesure')
#             )
#             .annotate(
#                 total_quantite=Sum('quantite')
#             )
#             .order_by('district_nom', 'composant_nom')
#         )
#
#         # Organiser par district
#         data = {}
#         for s in stocks:
#             district = s['district_nom']
#             if district not in data:
#                 data[district] = []
#             data[district].append({
#                 'composant': s['composant_nom'],
#                 'quantite': s['total_quantite'],
#                 'unite': s['unite']
#             })
#
#         context['stocks_par_district'] = data
#         return context
class StockDistrictView(LoginRequiredMixin, TemplateView):
    template_name = 'kits/stock_district.html'
    login_url = 'account_login'
    allowed_roles = ['National', 'Regional', 'District']
    redirect_view_if_denied = 'public_dashboard'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        user = self.request.user
        user_district = user.centre.district if user.roleemployee == 'DistrictSanitaire' else None

        # Filtres
        district_id = self.request.GET.get('district_id')
        kit_type_id = self.request.GET.get('kit_type_id')
        statut_stock = self.request.GET.get('statut_stock')

        # Base queryset
        stocks = Stock.objects.select_related(
            'composant',
            'centre',
            'centre__region__district'
        ).filter(centre__region__district__isnull=False)

        if district_id:
            stocks = stocks.filter(centre__region__district_id=district_id)
        elif user_district:
            stocks = stocks.filter(centre__region__district=user_district)

        if kit_type_id:
            stocks = stocks.filter(composant__kit_type_id=kit_type_id)

        if statut_stock:
            if statut_stock == 'critique':
                stocks = stocks.filter(quantite__lte=F('composant__seuil_alerte'))
            elif statut_stock == 'alerte':
                stocks = stocks.filter(
                    quantite__gt=F('composant__seuil_alerte'),
                    quantite__lte=F('composant__seuil_alerte') * 2
                )
            elif statut_stock == 'suffisant':
                stocks = stocks.filter(quantite__gt=F('composant__seuil_alerte') * 2)

        # Résumé par district et kit type
        summary = (
            stocks
            .values('centre__region__district__nom', 'composant__kit_type__nom')
            .annotate(
                total_quantite=Sum('quantite'),
                centres_count=models.Count('centre', distinct=True)
            )
            .order_by('centre__region__district__nom')
        )

        # Détails par composant
        details = (
            stocks
            .values(
                'centre__region__district__nom',
                'composant__nom',
                'composant__kit_type__nom',
                'composant__unite_mesure',
                'composant__seuil_alerte'
            )
            .annotate(
                total_quantite=Sum('quantite'),
                min_expiration=Min('date_expiration')
            )
            .order_by('centre__region__district__nom', 'composant__nom')
        )

        # Graphiques : quantités par district et par kit type
        districts_data = (
            stocks
            .values('centre__region__district__nom')
            .annotate(total=Sum('quantite'))
            .order_by('-total')
        )

        kit_types_data = (
            stocks
            .values('composant__kit_type__nom')
            .annotate(total=Sum('quantite'))
            .order_by('-total')
        )

        # Alerte stock bas ou expiration proche
        alertes = stocks.filter(
            Q(quantite__lte=F('composant__seuil_alerte')) |
            Q(date_expiration__lte=timezone.now().date() + timedelta(days=30))
        ).order_by('date_expiration')[:10]

        context.update({
            'summary': summary,
            'details': details,
            'districts_data': districts_data,
            'kit_types_data': kit_types_data,
            'alertes': alertes,

            'districts': DistrictSanitaire.objects.all().order_by('nom'),
            'kit_types': KitCategorie.objects.all().order_by('nom'),
            'statut_stock_choices': [
                ('all', 'Tous'),
                ('critique', 'Critique'),
                ('alerte', 'Alerte'),
                ('suffisant', 'Suffisant')
            ],

            'current_district': district_id,
            'current_kit_type': kit_type_id,
            'current_statut': statut_stock,

            'can_export': user.roleemployee in ['National', 'Regional'],
            'can_manage': user.roleemployee == 'National',
        })
        return context
